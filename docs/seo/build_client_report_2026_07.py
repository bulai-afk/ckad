#!/usr/bin/env python3
"""Клиентский Word-отчёт по центр-каталогизации.рф за июль 2026.

Структура как у образца opk.spb.ru. Цифры только из выгрузок Метрики,
уже лежащих в docs/seo (traffic-audit-2026-07.xlsx и
«Центр каталогизации отчет июнь-июль 2026.xlsx»). Срезы, которых
в пакете нет (география, устройства, год к году, Topvisor),
не выдумываются — в разделе остаётся пометка, что снять в следующем цикле.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "Отчет центр-каталогизации.рф за июль 2026.docx"
OUT_PATH_ASCII = ROOT / "ckad-seo-report-july-2026.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0x1A, 0x56, 0x8C)
MUTED = RGBColor(0x5A, 0x64, 0x72)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RED = RGBColor(0xB4, 0x23, 0x18)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1F3864"
ROW_ALT = "F4F7FA"
KPI_BG = "E8EEF5"

# Период в источнике: отчёт снят 30.07.2026, июль без 31-го числа.
PERIOD = "1 июля 2026 — 30 июля 2026"
COMPARE = "1 июня 2026 — 30 июня 2026"
COUNTER = "98614192"


def ru_pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%".replace(".", ",")


def rel_pct(new: float, old: float) -> str:
    if old == 0:
        return ""
    delta = (new - old) / old * 100
    if abs(delta) < 0.5:
        return ""
    return f"{delta:+.0f}%"


def pp_delta(new: float, old: float) -> str:
    """Дельта отказов в процентных пунктах, как в образце opk (26%-1)."""
    diff = round((new - old) * 100)
    if diff == 0:
        return "0"
    return f"{diff:+d}".replace("+", "+")


def depth_delta(new: float, old: float) -> str:
    return rel_pct(new, old)


def parse_hms(text: str) -> int:
    parts = [int(p) for p in text.split(":")]
    if len(parts) == 3:
        return parts[0] * 3600 + parts[1] * 60 + parts[2]
    return parts[0] * 60 + parts[1]


def fmt_mmss(seconds: int) -> str:
    minutes, sec = divmod(int(seconds), 60)
    hours, minutes = divmod(minutes, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{sec:02d}"
    return f"{minutes:02d}:{sec:02d}"


def time_delta(new_hms: str, old_hms: str) -> str:
    return rel_pct(parse_hms(new_hms), parse_hms(old_hms))


def visits_cell(new: float, old: float) -> str:
    return f"{int(new)}{rel_pct(new, old)}"


def bounce_cell(new: float, old: float) -> str:
    return f"{ru_pct(new, 1)}{pp_delta(new, old)}"


def depth_cell(new: float, old: float) -> str:
    return f"{new:.2f}{depth_delta(new, old)}".replace(".", ",")


def time_cell(new_hms: str, old_hms: str) -> str:
    return f"{fmt_mmss(parse_hms(new_hms))}{time_delta(new_hms, old_hms)}"


def set_run_font(run, *, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for key, value_cm in kwargs.items():
        node = OxmlElement(f"w:{key}")
        node.set(qn("w:w"), str(int(value_cm * 567)))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="C5CDD6", sz="4") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def prevent_row_split(row) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def write_cell(cell, text, *, bold=False, size=10, color=None, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    set_cell_margins(cell, top=0.08, bottom=0.08, left=0.12, right=0.12)
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = 1  # center


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=8, space_before=0, align="left", name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, name=name)
    return p


def add_heading_styled(doc, text, *, bookmark=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "1")
    pPr.append(outline)
    if bookmark:
        tag = OxmlElement("w:bookmarkStart")
        tag.set(qn("w:id"), bookmark)
        tag.set(qn("w:name"), bookmark)
        p._p.insert(0, tag)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark)
        p._p.append(end)
    return p


def add_caption(doc, text):
    return add_para(doc, text, size=9, color=MUTED, space_after=10, space_before=0)


def add_note(doc, text):
    return add_para(doc, text, size=10, color=MUTED, space_after=10, align="justify")


def kpi_row(doc, items: list[tuple[str, str, str | None]]):
    table = doc.add_table(rows=2, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="D5DCE4", sz="4")
    width = 16.5 / len(items)
    for i, (label, value, delta) in enumerate(items):
        cell0 = table.rows[0].cells[i]
        cell1 = table.rows[1].cells[i]
        write_cell(cell0, label, size=9, color=MUTED, align="center", fill=KPI_BG, bold=True)
        cell1.text = ""
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        set_run_font(run, size=20, bold=True, color=NAVY)
        if delta:
            p2 = cell1.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(6)
            sign_color = GREEN if delta.startswith("+") or "лучше" in delta else (
                RED if delta.startswith("-") else MUTED
            )
            # Снижение отказов — это хорошо.
            if "отказ" in label.lower() and delta.startswith("-"):
                sign_color = GREEN
            r2 = p2.add_run(delta)
            set_run_font(r2, size=10, bold=True, color=sign_color)
        shade(cell1, "FFFFFF")
        cell0.width = Cm(width)
        cell1.width = Cm(width)
        prevent_row_split(table.rows[0])
        prevent_row_split(table.rows[1])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def data_table(doc, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="C5CDD6", sz="4")
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, bold=True, size=9, color=WHITE, align="center", fill=HEADER_BG)
    for r_idx, row in enumerate(rows):
        fill = ROW_ALT if r_idx % 2 == 1 else "FFFFFF"
        prevent_row_split(table.rows[r_idx + 1])
        for c_idx, val in enumerate(row):
            align = "left" if c_idx in (0, 1) and not str(val).isdigit() else "center"
            if c_idx == 0:
                align = "center"
            write_cell(
                table.rows[r_idx + 1].cells[c_idx],
                val,
                size=9,
                align=align,
                fill=fill,
                bold=(r_idx == len(rows) - 1 and str(row[0]).upper().startswith("ИТОГО")),
            )
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("центр-каталогизации.рф  ·  отчёт о продвижении  ·  июль 2026")
    set_run_font(run, size=8, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Счётчик Яндекс.Метрики ")
    set_run_font(run, size=8, color=MUTED)
    run = fp.add_run(COUNTER)
    set_run_font(run, size=8, color=MUTED)
    run = fp.add_run("   ·   ")
    set_run_font(run, size=8, color=MUTED)

    # PAGE x of y
    run = fp.add_run("стр. ")
    set_run_font(run, size=8, color=MUTED)
    fld1 = OxmlElement("w:fldSimple")
    fld1.set(qn("w:instr"), "PAGE")
    r1 = OxmlElement("w:r")
    r1pr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    r1pr.append(sz)
    t1 = OxmlElement("w:t")
    t1.text = "1"
    r1.append(r1pr)
    r1.append(t1)
    fld1.append(r1)
    fp._p.append(fld1)
    run = fp.add_run(" из ")
    set_run_font(run, size=8, color=MUTED)
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES")
    r2 = OxmlElement("w:r")
    r2pr = OxmlElement("w:rPr")
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "16")
    r2pr.append(sz2)
    t2 = OxmlElement("w:t")
    t2.text = "1"
    r2.append(r2pr)
    r2.append(t2)
    fld2.append(r2)
    fp._p.append(fld2)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)
    add_header_footer(doc)

    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(11)
    styles.font.color.rgb = RGBColor(0x22, 0x26, 0x2A)

    # --- Обложка ---
    add_para(doc, "Отчет о продвижении сайта", size=14, color=ACCENT, space_after=4, align="center")
    add_para(doc, "центр-каталогизации.рф", size=26, bold=True, color=NAVY, space_after=14, align="center")
    add_para(
        doc,
        "Отчетный период с 1 июля 2026 по 30 июля 2026",
        size=12,
        bold=True,
        color=NAVY,
        space_after=2,
        align="center",
    )
    add_para(
        doc,
        "в сравнении с периодом с 1 июня 2026 по 30 июня 2026",
        size=12,
        color=MUTED,
        space_after=6,
        align="center",
    )
    add_para(
        doc,
        "Год к году (июль 2025) в исходных выгрузках Метрики не приложен — блок отмечен отдельно.",
        size=9,
        color=MUTED,
        space_after=18,
        align="center",
    )

    add_heading_styled(doc, "Содержание отчета")
    toc = [
        "Список проведенных работ",
        "Общая посещаемость и поведенческие факторы",
        "Переходы из поисковых систем",
        "Распределение трафика по каналам",
        "Тип устройств, с которых были переходы за отчетный период",
        "География посетителей",
        "Динамика поискового трафика",
        "Поисковый трафик за отчетный период в сравнении с аналогичным периодом за прошлый год",
        "Динамика переходов из поисковых систем за отчетный период",
        "Популярные посадочные страницы по всем источникам",
        "Популярные поисковые фразы",
        "Переходы с сайтов",
        "Популярные посадочные страницы из поисковых систем",
        "Источники цифр и что снять в следующем цикле",
    ]
    for i, item in enumerate(toc, 1):
        add_para(doc, f"{i}.  {item}", size=11, space_after=3, color=NAVY)

    add_caption(
        doc,
        "Цифры — Яндекс.Метрика, счётчик 98614192. Июль снят 30.07.2026 "
        "(1–30 июля), сравнение с 1–30 июня 2026. "
        "Это тот же счётчик на старой (июнь) и новой (июль) версии сайта.",
    )

    # --- Список работ ---
    add_heading_styled(doc, "Список проведенных работ")
    add_para(
        doc,
        "Задачи периода: закрепить новую структуру сайта после смены адресов, "
        "не потерять поисковый трафик со старых URL и собрать первую чистую "
        "картину Метрики по новой версии.",
        size=11,
        space_after=8,
        align="justify",
    )
    works = [
        "Сопровождение новой версии сайта: короткие адреса (/catalogization, /about, /contacts, услуги) вместо старых длинных путей /uslugi/, /o-kompanii/, /kontakty/.",
        "21 июля — уточнены 301-редиректы со старых URL раздела /uslugi на актуальные страницы услуг, чтобы переходы из поиска не упирались в 404.",
        "Доработка главной страницы и раздела статей (контент и блоки на новой вёрстке).",
        "Описание блока новостей вынесено в настройки раздела.",
        "Контроль единого счётчика Яндекс.Метрики 98614192 на обеих версиях сайта.",
        "Сбор отчётности: источники, страницы входа, поисковые фразы, цель «клик по номеру телефона»; сверка заявок с сайта.",
        "Семантика и мета-теги ключевых посадочных (каталогизация / ГОЗ / СФО / ФКП) — ядро ведётся в рабочем xlsx; в августе приоритет — title и description, где в выдаче много показов и мало кликов.",
    ]
    for i, item in enumerate(works, 1):
        add_para(doc, f"{i}. {item}", size=11, space_after=4, align="justify")

    add_para(doc, "Итоги месяца", size=12, bold=True, color=NAVY, space_before=8, space_after=6)
    add_para(
        doc,
        "Новая версия сильнее старой по объёму и качеству сессий: визиты 457 против 298 (+53%), "
        "отказы 18,6% против 25,5%. Доля поиска выросла примерно до 74%. "
        "Яндекс: 269 визитов (+74%), Google: 67 (+168%). "
        "Телефонные клики остались 3 = 3 при росте трафика — лиды не выросли пропорционально визитам. "
        "На новой версии зафиксированы 8 заявок с сайта (со старой версии сравнительной базы по заявкам нет). "
        "В выдаче уже есть и бренд, и коммерческие/информационные запросы; кликов по небренду пока мало — "
        "нужны правки сниппетов, а не новый объём «ради трафика».",
        size=11,
        space_after=8,
        align="justify",
    )
    add_para(doc, "Позиции", size=12, bold=True, color=NAVY, space_before=4, space_after=4)
    add_para(
        doc,
        "Отдельный трекер позиций (Topvisor и аналоги) в июле не вёлся — проценты «доля запросов в ТОП» "
        "по России и Санкт-Петербургу в этот отчёт не подставляем. "
        "Косвенно: органика выросла, бренд в поиске отрабатывает, коммерческие посадки в зоне средних позиций "
        "(ориентир Вебмастера ~6–9) при низком CTR на ряде URL. "
        "Для следующего отчёта нужен снимок ТОП по гео «Россия» и «Санкт-Петербург», Яндекс и Google.",
        size=11,
        space_after=8,
        align="justify",
    )
    add_para(doc, "Конверсии", size=12, bold=True, color=NAVY, space_before=4, space_after=4)
    add_para(
        doc,
        "Сегмент «только поисковый трафик» по целям в пакете выгрузок не приложен. "
        "Ниже — цели по всему трафику сайта, июль vs июнь:",
        size=11,
        space_after=6,
        align="justify",
    )
    data_table(
        doc,
        ["Цель", "Июнь 2026", "Июль 2026", "Комментарий"],
        [
            ["Клик по номеру телефона", "3 целевых визита", "3 целевых визита", "Абсолют не вырос при +53% визитов"],
            ["CR телефон от всех визитов", "~1,0%", "~0,66%", "Конверсионность трафика ниже"],
            ["Заявки с сайта", "нет данных", "8", "База для следующих месяцев; со старой не сравниваем"],
            ["Форма / email / мессенджер детально", "нет выгрузки", "нет выгрузки", "Снять цель × страница входа"],
        ],
        col_widths=[4.8, 3.4, 3.4, 5.0],
    )

    # --- Сводка ---
    add_heading_styled(doc, "Общая посещаемость и поведенческие факторы")
    add_caption(doc, "Яндекс.Метрика → Источники, сводка · визиты · 01.07–30.07.2026 vs 01.06–30.06.2026")
    visits_jul, visits_jun = 457, 298
    users_jul, users_jun = 238, 187
    depth_jul, depth_jun = 1.663, 1.721
    views_jul = round(visits_jul * depth_jul)  # 760
    views_jun = round(visits_jun * depth_jun)  # 513
    kpi_row(
        doc,
        [
            ("Количество посетителей", str(users_jul), rel_pct(users_jul, users_jun)),
            ("Количество визитов", str(visits_jul), rel_pct(visits_jul, visits_jun)),
            ("Количество просмотров", str(views_jul), rel_pct(views_jul, views_jun)),
        ],
    )
    add_caption(doc, "Просмотры посчитаны как визиты × глубина (в сводке Метрики отдельной колонкой не выгружались).")
    kpi_row(
        doc,
        [
            ("Глубина просмотра", f"{depth_jul:.2f}".replace(".", ","), rel_pct(depth_jul, depth_jun)),
            ("Длительность визита", "02:33", time_delta("00:02:33", "00:02:45")),
            ("Показатель отказов", ru_pct(0.186), "лучше на 6,9 п.п."),
        ],
    )

    # --- Поиск ---
    add_heading_styled(doc, "Переходы из поисковых систем")
    add_caption(doc, "Яндекс.Метрика → Источники, сводка · поисковые системы")
    kpi_row(
        doc,
        [
            ("Яндекс", "269", rel_pct(269, 155)),
            ("Google", "67", rel_pct(67, 25)),
            ("Bing", "2", rel_pct(2, 4)),
        ],
    )
    add_note(
        doc,
        "Итого поиск: 338 визитов в июле против 184 в июне (+84%). "
        "Доля органики в визитах сайта — около 74% (в июне около 62%). "
        "Bing — 2 визита со 100% отказов, в оценке SEO не учитываем.",
    )

    # --- Каналы ---
    add_heading_styled(doc, "Распределение трафика по каналам")
    kpi_row(
        doc,
        [
            ("Переходы из поисковых систем", "338", rel_pct(338, 184)),
            ("Прямые заходы", "106", rel_pct(106, 70)),
            ("Переходы по рекламе", "8", rel_pct(8, 21)),
        ],
    )
    data_table(
        doc,
        ["№", "Источник", "Визиты", "Отказы", "Глубина просмотра", "Время на сайте"],
        [
            ["1", "Переходы из поисковых систем", visits_cell(338, 184), bounce_cell(0.118, 0.125), depth_cell(1.53, 1.68), time_cell("00:02:28", "00:02:09")],
            ["2", "Прямые заходы", visits_cell(106, 70), bounce_cell(0.4151, 0.7), depth_cell(2.132, 1.443), time_cell("00:02:56", "00:01:00")],
            ["3", "Переходы по рекламе (Яндекс.Директ)", visits_cell(8, 21), bounce_cell(0.125, 0.0), depth_cell(1.5, 1.0), time_cell("00:01:28", "00:00:48")],
            ["4", "Переходы по ссылкам на сайтах", visits_cell(5, 8), bounce_cell(0.2, 0.125), "1", "00:17"],
            ["5", "Внутренние переходы", visits_cell(0, 10), "—", "—", "—"],
            ["", "Итого", visits_cell(457, 298), bounce_cell(0.186, 0.255), depth_cell(1.663, 1.721), time_cell("00:02:33", "00:02:45")],
        ],
        col_widths=[1.0, 5.4, 2.4, 2.3, 2.8, 2.6],
    )
    add_note(
        doc,
        "Прямые заходы выросли на 51%, но отказы по ним 41,5% — смесь бренда и сессий без referrer. "
        "Директ сжался до 8 визитов (−62%) и с оценкой SEO не смешиваем. "
        "Внутренние переходы в июле обнулились (в июне было 10) — проверить фильтры счётчика, на вывод по поиску не влияет.",
    )

    # --- Устройства ---
    add_heading_styled(doc, "Тип устройств, с которых были переходы за отчетный период")
    add_note(
        doc,
        "В пакете выгрузок за июль 2026 отчёта «Технологии → Устройства» нет. "
        "Таблица ПК / смартфоны / планшеты в этот файл не выдумана. "
        "Для августа: Метрика → Технологии → Устройства, тот же период и сравнение с предыдущим месяцем, "
        "колонки визиты / отказы / глубина / время.",
    )

    # --- География ---
    add_heading_styled(doc, "География посетителей")
    add_note(
        doc,
        "В пакете выгрузок отчёта «Аудитория → География» нет. "
        "Топ регионов с дельтами к июню в этот файл не выдуман. "
        "Для августа: Метрика → Аудитория → География, топ-10, визиты / отказы / глубина / время, "
        "на карточках — топ-6 плюс «Другие».",
    )

    # --- Динамика поиска vs июнь ---
    add_heading_styled(doc, "Динамика поискового трафика")
    add_caption(doc, "Сегмент «Переходы из поисковых систем» · июль 2026 vs июнь 2026")
    kpi_row(
        doc,
        [
            ("Визиты", "338", rel_pct(338, 184)),
            ("Отказы", ru_pct(0.118), "↓ к 12,5% в июне"),
            ("Глубина просмотров", "1,53", rel_pct(1.53, 1.68)),
            ("Время на сайте", "02:28", time_delta("00:02:28", "00:02:09")),
        ],
    )
    add_note(
        doc,
        "Посуточный график в выгрузку не входил — в образце opk он стоит скрином из Метрики. "
        "Итоги периода: поиск вырос с 184 до 338 визитов. "
        "Отказы органики низкие (Яндекс 11,9%, Google 7,5%). "
        "Глубина органики ~1,5 — зона роста (перелинковка и CTA на услугах).",
    )

    # --- YoY ---
    add_heading_styled(doc, "Поисковый трафик за отчетный период в сравнении с аналогичным периодом за прошлый год")
    add_note(
        doc,
        "Сравнение с 1–31 июля 2025 в выгрузках Метрики не приложено. "
        "Старый PDF по сайту за 11.06.2025–10.07.2025 (182 посетителя, 257 визитов) "
        "к календарному июлю 2025 не приводится: период другой, как год к году не используем. "
        "Для следующего отчёта: тот же сегмент поиска, группировка по дням, "
        "второй период — 01.07.2025–31.07.2025.",
    )

    # --- По системам таблица ---
    add_heading_styled(doc, "Динамика переходов из поисковых систем за отчетный период")
    data_table(
        doc,
        ["№", "Поисковая система", "Визиты", "Отказы", "Глубина просмотра", "Время на сайте"],
        [
            ["1", "Яндекс", visits_cell(269, 155), bounce_cell(0.119, 0.1097), depth_cell(1.543, 1.645), time_cell("00:02:25", "00:01:41")],
            ["2", "Google", visits_cell(67, 25), bounce_cell(0.0746, 0.12), depth_cell(1.493, 1.92), time_cell("00:02:48", "00:04:21")],
            ["3", "Bing", visits_cell(2, 4), bounce_cell(1.0, 0.75), depth_cell(1.0, 1.5), time_cell("00:00:00", "00:06:12")],
        ],
        col_widths=[1.0, 4.2, 2.6, 2.4, 3.1, 3.2],
    )
    add_note(
        doc,
        "Яндекс — основной канал (59% всех визитов сайта). "
        "Google вырос с низкой базы, отказы лучше (7,5%). "
        "Время в Google снизилось (04:21 → 02:48) при росте объёма — смотреть посадки, не только средний визит.",
    )

    # --- Посадки все источники ---
    add_heading_styled(doc, "Популярные посадочные страницы по всем источникам")
    add_caption(doc, "Яндекс.Метрика → страницы входа · июль vs июнь 2026 · топ-10 новой структуры")
    data_table(
        doc,
        ["№", "URL", "Визитов"],
        [
            ["1", "https://центр-каталогизации.рф/", visits_cell(281, 174)],
            ["2", "https://центр-каталогизации.рф/catalogization", visits_cell(37, 4)],
            ["3", "https://центр-каталогизации.рф/catalogization/katalogizatsiya-produktsii-po-goz-pod-klyuch", visits_cell(20, 2)],
            ["4", "https://центр-каталогизации.рф/about", visits_cell(20, 1)],
            ["5", "https://центр-каталогизации.рф/training-center/seminary-po-katalogizatsii", visits_cell(16, 4)],
            ["6", "https://центр-каталогизации.рф/articles/zachem-nuzhna-katalogizatsiya-produktsii-i-chto-eto-takoe", visits_cell(12, 3)],
            ["7", "https://центр-каталогизации.рф/catalogization/razrabotka-standartnyh-formatov-opisaniya-sfo", visits_cell(12, 0)],
            ["8", "https://центр-каталогизации.рф/articles", visits_cell(10, 2)],
            ["9", "https://центр-каталогизации.рф/other-services/dokumenty-dlya-primeneniya-inostrannoy-produktsii-v-vvst", visits_cell(7, 0)],
            ["10", "https://центр-каталогизации.рф/catalogization/proverka-nalichiya-produktsii-v-fkp", visits_cell(5, 1)],
        ],
        col_widths=[1.0, 12.6, 3.0],
    )
    add_note(
        doc,
        "Главная забирает 281 из 457 визитов (~61%). Хаб /catalogization вырос с 4 до 37. "
        "Лучшая коммерческая посадка — каталогизация по ГОЗ (20 визитов, глубина 2,45, отказы 10%). "
        "Старые URL июня (/uslugi/…, /o-kompanii/, старая статья) в июле как входы = 0: миграция на редиректах видна на одном счётчике.",
    )

    # --- Фразы ---
    add_heading_styled(doc, "Популярные поисковые фразы")
    add_caption(
        doc,
        "Яндекс.Метрика → поисковые фразы, июль 2026. "
        "В выгрузке — 216 визитов с известной фразой; построчный топ-10 отдельных запросов "
        "в пакет не входил, ниже кластеры (как они собраны в рабочем отчёте).",
    )
    data_table(
        doc,
        ["№", "Фраза / кластер", "Визитов", "Доля от 216"],
        [
            ["1", "Бренд: «центр каталогизации», ООО, домен", "112", "52%"],
            ["2", "Коммерция + info (небренд)", "62", "29%"],
            ["3", "Телефон / ИНН / email / контакты", "26", "12%"],
            ["4", "Шум «каталог сопровождение» (не цель)", "16", "7%"],
        ],
        col_widths=[1.0, 9.5, 2.8, 3.2],
    )
    add_note(
        doc,
        "Рост июля в значительной степени брендовый. "
        "Небренд — зона масштабирования: каталогизация / что это / зачем; каталогизация продукции по ГОЗ; "
        "СФО / ФКП / ФНН; семинары; иностранная продукция ВВСТ. "
        "Запросы «каталог сопровождение» не продвигаем. "
        "Для августа нужна сырая таблица топ-10 фраз из Метрики (фраза + визиты + % к июню).",
    )

    # --- Рефералы ---
    add_heading_styled(doc, "Переходы с сайтов")
    data_table(
        doc,
        ["№", "Источник", "Визиты"],
        [
            ["1", "https://checko.ru/", visits_cell(2, 2)],
            ["2", "https://ekspertvvt.bitrix24.ru/", visits_cell(2, 0)],
            ["3", "https://metrika.yandex.ru/", visits_cell(1, 0)],
            ["4", "https://hosting.timeweb.ru/", visits_cell(0, 6)],
        ],
        col_widths=[1.0, 12.6, 3.0],
    )
    add_note(
        doc,
        "Внешний ссылочный трафик почти отсутствует (~1% визитов). "
        "checko.ru — каталог компаний без глубины. Bitrix24 и Метрика — служебные, не маркетинг. "
        "Хостинг Timeweb давал 6 визитов в июне и 0 в июле. "
        "Как в образце opk, имеет смысл планомерно добавлять отраслевые каталоги и партнёрские упоминания.",
    )

    # --- Посадки из поиска ---
    add_heading_styled(doc, "Популярные посадочные страницы из поисковых систем")
    add_note(
        doc,
        "Отдельная выгрузка «страницы входа» с сегментом только поиск в пакете не приложена. "
        "Доля поиска в июле ~74%, поэтому верх таблицы по всем источникам в основном поисковый; "
        "главная дополнительно принимает прямые заходы. "
        "Ниже — соответствие кластеров спроса посадочным новой структуры (для работы, не как замена сегмента Метрики).",
    )
    data_table(
        doc,
        ["Кластер", "Целевая страница", "Что видно в июле"],
        [
            ["Бренд", "/", "281 вход; бренд в выдаче отрабатывает"],
            ["Каталогизация / что это / зачем", "/articles/zachem-nuzhna-katalogizatsiya-produktsii-i-chto-eto-takoe", "12 входов; в выдаче много показов, CTR низкий"],
            ["Каталогизация продукции по ГОЗ", "/catalogization/katalogizatsiya-produktsii-po-goz-pod-klyuch", "20 входов, глубина 2,45 — лучшая коммерция"],
            ["СФО / ФКП / ФНН", "/catalogization/…-sfo, проверка ФКП, КИ-МВН", "точечные входы; проверка ФКП — сильный CTR в Вебмастере"],
            ["Семинар / курсы / обучение", "/training-center/seminary-po-katalogizatsii", "16 входов, отказы 6%"],
            ["Иностранная продукция / ВВСТ", "/other-services/dokumenty-…-vvst", "7 входов, long-tail"],
        ],
        col_widths=[4.4, 6.6, 5.6],
    )

    # --- Источники ---
    add_heading_styled(doc, "Источники цифр и что снять в следующем цикле")
    add_para(
        doc,
        "В этом файле нет выдуманных строк Метрики. Основа — выгрузки, уже разобранные во внутренних xlsx "
        "(traffic-audit-2026-07.xlsx, «Центр каталогизации отчет июнь-июль 2026.xlsx»), "
        "снятые 30.07.2026.",
        size=11,
        space_after=6,
        align="justify",
    )
    data_table(
        doc,
        ["Срез", "Статус в этом отчёте", "Что прислать в августе"],
        [
            ["Сводка, каналы, поисковые системы", "Есть", "Тот же отчёт за полный календарный месяц"],
            ["Страницы входа", "Есть (все источники)", "Плюс сегмент только поиск, топ-10"],
            ["Кластеры фраз", "Есть", "Сырой топ-10 фраз с дельтой к прошлому месяцу"],
            ["Рефералы", "Есть", "Как есть"],
            ["Конверсии", "Телефон 3/3, заявки 8", "Цели × поиск и цели × страница входа"],
            ["Устройства", "Нет выгрузки", "Технологии → Устройства"],
            ["География", "Нет выгрузки", "Аудитория → География, топ-10"],
            ["График поиска vs прошлый месяц", "Только итоги", "Скрин по дням"],
            ["Поиск vs июль 2025", "Нет", "Тот же сегмент, второй период — июль 2025"],
            ["Позиции РФ / СПб", "Нет трекера", "Topvisor или аналог, Яндекс и Google"],
        ],
        col_widths=[5.2, 5.2, 6.2],
    )

    add_para(
        doc,
        "Отчёт собран по структуре образца opk.spb.ru (список работ + блоки Метрики). "
        "Блоки Google Search Console, ИКС и история индексирования сознательно не включались.",
        size=9,
        color=MUTED,
        space_before=8,
        space_after=0,
    )

    doc.save(OUT_PATH)
    doc.save(OUT_PATH_ASCII)
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(path)
