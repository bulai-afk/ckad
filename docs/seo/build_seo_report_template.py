#!/usr/bin/env python3
"""Пустой Word-шаблон ежемесячного SEO-отчёта.

Четыре периода: отчётный месяц, предыдущий, позапрошлый, тот же месяц год назад.
Ячейки с «……» заполняются вручную. В блоки графиков вставляются скрины Метрики.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT_ASCII = ROOT / "ckad-seo-report-template.docx"
OUT_RU = ROOT / "Шаблон SEO-отчета.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0x1A, 0x56, 0x8C)
MUTED = RGBColor(0x5A, 0x64, 0x72)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PLACE = RGBColor(0x9A, 0xA3, 0xAD)
HEADER_BG = "1F3864"
ROW_ALT = "F4F7FA"
KPI_BG = "E8EEF5"
CHART_BG = "F7F5EE"
BLANK = "……"

PERIODS = ["Отчётный", "Предыдущий", "Позапрошлый", "Год назад"]
DELTAS = ["Δ к пред.", "Δ к позапр.", "Δ к году"]


def set_run_font(run, *, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, **kwargs) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for key, value_cm in kwargs.items():
        node = OxmlElement(f"w:{key}")
        node.set(qn("w:w"), str(int(value_cm * 567)))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="C5CDD6", sz="4") -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def set_row_height(row, twips: int) -> None:
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:trHeight")
    el.set(qn("w:val"), str(twips))
    el.set(qn("w:hRule"), "atLeast")
    trPr.append(el)


def write_cell(cell, text, *, bold=False, size=9, color=None, align="center", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    is_blank = str(text).strip() in {BLANK, ""}
    set_run_font(run, size=size, bold=bold, color=PLACE if is_blank and color is None else color)
    set_cell_margins(cell, top=0.08, bottom=0.08, left=0.1, right=0.1)
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = 1


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=8, space_before=0, align="left"):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=15, bold=True, color=NAVY)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "1")
    p._p.get_or_add_pPr().append(outline)
    return p


def add_caption(doc, text):
    return add_para(doc, text, size=9, color=MUTED, space_after=8)


def data_table(doc, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, bold=True, size=8, color=WHITE, fill=HEADER_BG)
    for r_idx, row in enumerate(rows):
        fill = ROW_ALT if r_idx % 2 else "FFFFFF"
        for c_idx, val in enumerate(row):
            align = "left" if c_idx == 0 else "center"
            write_cell(table.rows[r_idx + 1].cells[c_idx], val, size=8, align=align, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def blank_row(first: str, n: int) -> list[str]:
    return [first] + [BLANK] * n


def compare_headers(first: str) -> list[str]:
    return [first] + PERIODS + DELTAS


def compare_widths(first_w: float) -> list[float]:
    rest = (16.6 - first_w) / 7
    return [first_w] + [rest] * 7


def chart_box(doc, title: str, hint: str) -> None:
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="C4B896", sz="8")
    cell = table.rows[0].cells[0]
    shade(cell, CHART_BG)
    set_row_height(table.rows[0], 2200)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Вставьте скриншот графика из Метрики")
    set_run_font(run, size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    run2 = p2.add_run(hint)
    set_run_font(run2, size=9, color=MUTED)
    cell.width = Cm(16.6)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("SEO-отчёт  ·  шаблон сравнительного анализа  ·  заполнить и сохранить копию")
    set_run_font(run, size=8, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("стр. ")
    set_run_font(run, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    fp._p.append(fld)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    add_header_footer(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # --- Как заполнять ---
    add_para(doc, "Шаблон сравнительного SEO-отчёта", size=20, bold=True, color=NAVY, align="center", space_after=4)
    add_para(doc, "центр-каталогизации.рф  ·  Яндекс.Метрика 98614192", size=12, color=ACCENT, align="center", space_after=12)

    add_heading(doc, "Как заполнять")
    add_para(
        doc,
        "1. Сохраните копию файла (Файл → Сохранить как) с именем месяца, например «Отчет … за август 2026». "
        "Этот шаблон не затирайте.",
        size=11,
        space_after=4,
        align="justify",
    )
    add_para(
        doc,
        "2. Серые «……» замените цифрами из Метрики. Δ считайте так: "
        "(отчётный − сравнение) / сравнение × 100%. Для отказов удобнее пункты: отчёт% − сравнение%.",
        size=11,
        space_after=4,
        align="justify",
    )
    add_para(
        doc,
        "3. Графики: в жёлтые рамки вставьте скрин (ПКМ → Вставить). "
        "В Метрике сравнение только двух периодов: сначала отчётный vs предыдущий, потом отчётный vs год назад.",
        size=11,
        space_after=8,
        align="justify",
    )
    data_table(
        doc,
        ["Колонка", "Что писать", "Пример для отчёта за август 2026"],
        [
            ["Отчётный", "месяц, который сдаём клиенту", "01.08.2026 – 31.08.2026"],
            ["Предыдущий", "прошлый календарный месяц", "01.07.2026 – 31.07.2026"],
            ["Позапрошлый", "ещё на месяц раньше", "01.06.2026 – 30.06.2026"],
            ["Год назад", "тот же календарный месяц год назад", "01.08.2025 – 31.08.2025"],
        ],
        col_widths=[3.2, 5.6, 7.8],
    )
    add_caption(
        doc,
        "Откуда брать: Сводка; Источники → Сводка / Поисковые системы / Фразы / Сайты; "
        "Технологии → Устройства; Аудитория → География; Содержание → Посадочные; Конверсии. "
        "Группировка по визитам. GSC, ИКС, индексирование в этот шаблон не входят.",
    )

    # --- Обложка отчёта ---
    add_heading(doc, "Обложка")
    add_para(doc, "Отчет о продвижении сайта", size=13, color=ACCENT, align="center", space_after=2)
    add_para(doc, "центр-каталогизации.рф", size=22, bold=True, color=NAVY, align="center", space_after=10)
    data_table(
        doc,
        ["Поле", "Заполнить"],
        [
            ["Отчётный период", "……  —  ……"],
            ["Предыдущий месяц", "……  —  ……"],
            ["Позапрошлый месяц", "……  —  ……"],
            ["Тот же период год назад", "……  —  ……"],
            ["Дата сборки отчёта", "……"],
        ],
        col_widths=[5.5, 11.1],
    )

    add_heading(doc, "Содержание отчёта")
    toc = [
        "Список проведенных работ",
        "Рамка периодов (4 месяца + дельты)",
        "Общая посещаемость и поведенческие факторы",
        "Переходы из поисковых систем",
        "Распределение трафика по каналам",
        "Тип устройств",
        "География посетителей",
        "Динамика поискового трафика (графики)",
        "Поисковый трафик vs прошлый год (график)",
        "Динамика переходов из поисковых систем",
        "Популярные посадочные по всем источникам",
        "Популярные поисковые фразы",
        "Переходы с сайтов",
        "Популярные посадочные из поиска",
        "Конверсии",
        "Позиции (если есть трекер)",
    ]
    for i, item in enumerate(toc, 1):
        add_para(doc, f"{i}.  {item}", size=11, space_after=2, color=NAVY)

    # --- Работы ---
    add_heading(doc, "Список проведенных работ")
    add_caption(doc, "Что сделали за отчётный месяц. Ненужные строки удалите, недостающие добавьте.")
    data_table(
        doc,
        ["№", "Работа", "URL / комментарий"],
        [[str(i), BLANK, BLANK] for i in range(1, 9)],
        col_widths=[1.2, 8.8, 6.6],
    )
    add_para(doc, "Задачи периода", size=12, bold=True, color=NAVY, space_after=4)
    add_para(doc, BLANK, size=11, color=PLACE, space_after=8)
    add_para(doc, "Итоги месяца (2–4 предложения)", size=12, bold=True, color=NAVY, space_after=4)
    add_para(doc, BLANK, size=11, color=PLACE, space_after=10)

    # --- Рамка ---
    add_heading(doc, "Рамка периодов")
    add_caption(doc, "Сводка сайта. Сначала заполните четыре колонки факта, потом три дельты.")
    metrics = [
        "Посетители",
        "Визиты",
        "Просмотры",
        "Отказы, %",
        "Глубина",
        "Время на сайте (мм:сс)",
        "Яндекс, визиты",
        "Google, визиты",
        "Поиск всего, визиты",
        "Доля поиска, %",
        "Прямые заходы",
        "Реклама",
        "Рефералы",
    ]
    data_table(
        doc,
        compare_headers("Показатель"),
        [blank_row(m, 7) for m in metrics],
        compare_widths(4.0),
    )

    # --- Посещаемость ---
    add_heading(doc, "Общая посещаемость и поведенческие факторы")
    add_caption(doc, "Метрика → Сводка. Карточки: отчётный месяц и Δ к предыдущему.")
    data_table(
        doc,
        ["Метрика", "Отчётный", "Δ к пред.", "Δ к позапр.", "Δ к году"],
        [blank_row(m, 4) for m in ["Посетители", "Визиты", "Просмотры", "Глубина", "Время", "Отказы"]],
        col_widths=[4.0, 3.15, 3.15, 3.15, 3.15],
    )
    chart_box(
        doc,
        "График 1. Посетители / визиты по дням — отчётный vs предыдущий",
        "Метрика → Посетители → группировка по дням → сравнение с предыдущим месяцем",
    )

    # --- Поиск ---
    add_heading(doc, "Переходы из поисковых систем")
    add_caption(doc, "Источники → Поисковые системы.")
    data_table(
        doc,
        compare_headers("Система"),
        [blank_row(s, 7) for s in ["Яндекс", "Google", "Bing", "Другие", "Итого поиск"]],
        compare_widths(3.2),
    )
    chart_box(
        doc,
        "График 2. Поиск по дням — отчётный vs предыдущий",
        "Сегмент «Переходы из поисковых систем», группировка по дням, vs предыдущий месяц",
    )

    # --- Каналы ---
    add_heading(doc, "Распределение трафика по каналам")
    add_caption(doc, "Источники → Сводка. Визиты и поведение — за отчётный месяц; Δ визитов — к трём сравнениям.")
    data_table(
        doc,
        ["№", "Источник", "Визиты отч.", "Δ пред.", "Δ позапр.", "Δ году", "Отказы", "Глубина", "Время"],
        [
            blank_row(s, 8)
            for s in [
                "1  Поиск",
                "2  Прямые",
                "3  Реклама",
                "4  Сайты (рефералы)",
                "5  Внутренние",
                "6  Соцсети / прочее",
                "Итого",
            ]
        ],
        col_widths=[0.8, 3.6, 1.8, 1.6, 1.7, 1.5, 1.6, 1.7, 1.5],
    )
    chart_box(
        doc,
        "График 3. Каналы отчётного месяца",
        "Источники → Сводка: круговая или столбцы по визитам",
    )

    # --- Устройства ---
    add_heading(doc, "Тип устройств")
    add_caption(doc, "Технологии → Устройства.")
    data_table(
        doc,
        compare_headers("Устройство"),
        [blank_row(s, 7) for s in ["ПК", "Смартфоны", "Планшеты"]],
        compare_widths(3.2),
    )
    data_table(
        doc,
        ["Устройство", "Визиты отч.", "Отказы", "Глубина", "Время"],
        [blank_row(s, 4) for s in ["ПК", "Смартфоны", "Планшеты"]],
        col_widths=[3.4, 3.3, 3.3, 3.3, 3.3],
    )
    chart_box(doc, "График 4. Устройства", "Технологии → Устройства, отчётный месяц")

    # --- Гео ---
    add_heading(doc, "География посетителей")
    add_caption(doc, "Аудитория → География, топ-10. Δ — по визитам.")
    data_table(
        doc,
        ["№", "Регион", "Отчётный", "Предыдущий", "Позапрошлый", "Год назад", "Отказы отч.", "Глубина", "Время"],
        [[str(i)] + [BLANK] * 8 for i in range(1, 11)],
        col_widths=[0.8, 3.6, 1.8, 2.0, 2.1, 1.8, 1.6, 1.5, 1.4],
    )
    chart_box(doc, "График 5. География", "Аудитория → География, отчётный месяц (карта или столбцы топ-6)")

    # --- Динамика поиска ---
    add_heading(doc, "Динамика поискового трафика")
    add_caption(doc, "Итоги сегмента «Поиск»: отчётный vs предыдущий (как карточки в образце opk).")
    data_table(
        doc,
        ["", "Визиты", "Отказы", "Глубина", "Время"],
        [
            ["Отчётный", BLANK, BLANK, BLANK, BLANK],
            ["Предыдущий", BLANK, BLANK, BLANK, BLANK],
            ["Δ", BLANK, BLANK, BLANK, BLANK],
        ],
        col_widths=[3.4, 3.3, 3.3, 3.3, 3.3],
    )

    add_heading(doc, "Поисковый трафик vs прошлый год")
    data_table(
        doc,
        ["", "Визиты", "Отказы", "Глубина", "Время"],
        [
            ["Отчётный", BLANK, BLANK, BLANK, BLANK],
            ["Год назад", BLANK, BLANK, BLANK, BLANK],
            ["Δ", BLANK, BLANK, BLANK, BLANK],
        ],
        col_widths=[3.4, 3.3, 3.3, 3.3, 3.3],
    )
    chart_box(
        doc,
        "График 6. Поиск по дням — отчётный vs тот же месяц год назад",
        "Сегмент «Переходы из поисковых систем», по дням, сравнение с 01.мм.гггг–31.мм.(год−1)",
    )

    # --- ПС детально ---
    add_heading(doc, "Динамика переходов из поисковых систем")
    add_caption(doc, "Отчётный месяц: визиты с дельтой к предыдущему + поведение.")
    data_table(
        doc,
        ["№", "Система", "Визиты + Δ пред.", "Отказы", "Глубина", "Время"],
        [[str(i), BLANK, BLANK, BLANK, BLANK, BLANK] for i in range(1, 6)],
        col_widths=[1.0, 3.6, 3.4, 2.8, 2.9, 2.9],
    )

    # --- Посадки все ---
    add_heading(doc, "Популярные посадочные страницы по всем источникам")
    add_caption(doc, "Содержание → Посадочные / страницы входа, топ-10. Δ — по визитам.")
    data_table(
        doc,
        ["№", "URL", "Отчётный", "Предыдущий", "Позапрошлый", "Год назад", "Δ пред."],
        [[str(i), BLANK, BLANK, BLANK, BLANK, BLANK, BLANK] for i in range(1, 11)],
        col_widths=[0.8, 6.4, 1.9, 2.1, 2.1, 1.8, 1.5],
    )

    # --- Фразы ---
    add_heading(doc, "Популярные поисковые фразы")
    add_caption(doc, "Источники → Поисковые фразы, топ-10.")
    data_table(
        doc,
        ["№", "Фраза", "Отчётный", "Предыдущий", "Позапрошлый", "Год назад", "Δ пред."],
        [[str(i), BLANK, BLANK, BLANK, BLANK, BLANK, BLANK] for i in range(1, 11)],
        col_widths=[0.8, 6.4, 1.9, 2.1, 2.1, 1.8, 1.5],
    )

    # --- Рефералы ---
    add_heading(doc, "Переходы с сайтов")
    add_caption(doc, "Источники → Сайты.")
    data_table(
        doc,
        ["№", "Источник", "Отчётный", "Предыдущий", "Позапрошлый", "Год назад"],
        [[str(i), BLANK, BLANK, BLANK, BLANK, BLANK] for i in range(1, 9)],
        col_widths=[0.9, 6.5, 2.3, 2.3, 2.3, 2.3],
    )

    # --- Посадки поиск ---
    add_heading(doc, "Популярные посадочные страницы из поисковых систем")
    add_caption(doc, "Те же посадочные, сегмент только «Переходы из поисковых систем».")
    data_table(
        doc,
        ["№", "URL", "Отчётный", "Предыдущий", "Позапрошлый", "Год назад", "Δ пред."],
        [[str(i), BLANK, BLANK, BLANK, BLANK, BLANK, BLANK] for i in range(1, 11)],
        col_widths=[0.8, 6.4, 1.9, 2.1, 2.1, 1.8, 1.5],
    )

    # --- Конверсии ---
    add_heading(doc, "Конверсии")
    add_caption(doc, "Метрика → Конверсии. Если есть — продублируйте строки сегментом «только поиск».")
    data_table(
        doc,
        compare_headers("Цель"),
        [
            blank_row(s, 7)
            for s in [
                "Клик по телефону",
                "CR телефон, %",
                "Заявки с сайта",
                "Форма / email / мессенджер",
                "…… (добавьте цель)",
            ]
        ],
        compare_widths(4.2),
    )

    # --- Позиции ---
    add_heading(doc, "Позиции")
    add_caption(doc, "Topvisor или аналог на последний день отчётного месяца. Нет трекера — напишите «нет данных».")
    data_table(
        doc,
        ["Гео / ПС", "Отчётный, % в ТОП", "Предыдущий", "Позапрошлый", "Год назад"],
        [
            blank_row(s, 4)
            for s in [
                "Россия · Яндекс",
                "Россия · Google",
                "Санкт-Петербург · Яндекс",
                "Санкт-Петербург · Google",
            ]
        ],
        col_widths=[5.0, 3.0, 2.9, 2.9, 2.8],
    )

    add_para(
        doc,
        "Шаблон по структуре образца opk.spb.ru. Пересобрать пустой файл: python3 docs/seo/build_seo_report_template.py",
        size=8,
        color=MUTED,
        space_before=12,
        space_after=0,
    )

    doc.save(OUT_ASCII)
    doc.save(OUT_RU)
    return OUT_ASCII


if __name__ == "__main__":
    print(build())
    print(OUT_RU)
